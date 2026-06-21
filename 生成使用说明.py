"""生成对话画布使用说明 Word 文档"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# 样式设置
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5

# ═══════════════════════════════════════════════════
#  标题
# ═══════════════════════════════════════════════════
title = doc.add_heading('对话画布 Conversation Canvas', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('使用说明 v0.5.0')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph()

# ═══════════════════════════════════════════════════
#  1. 这是什么
# ═══════════════════════════════════════════════════
doc.add_heading('1. 这是什么', level=1)

doc.add_paragraph(
    '对话画布是一个可视化 AI 对话工具。和普通的聊天界面不同，它把对话变成节点图——'
    '每个问题和回答都是一个节点，你可以看到整个对话的结构。'
)
doc.add_paragraph(
    '更重要的是，它支持 Agent 模式：AI 可以自动帮你创建文件、执行命令、搜索内容，'
    '并且在节点图上实时展示它的每一步操作。'
)

doc.add_heading('两种模式', level=2)

table = doc.add_table(rows=3, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['模式', '用途', '适合场景']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h

data = [
    ['对话模式', '和 AI 聊天，分叉对话，对比回复', '日常问答、头脑风暴'],
    ['Agent 模式', 'AI 自动调用工具完成任务', '创建文件、写代码、批量操作'],
]
for r, row_data in enumerate(data):
    for c, val in enumerate(row_data):
        table.rows[r+1].cells[c].text = val

doc.add_paragraph()

# ═══════════════════════════════════════════════════
#  2. 安装
# ═══════════════════════════════════════════════════
doc.add_heading('2. 安装', level=1)

doc.add_paragraph('这是一个单文件应用，不需要安装任何依赖。')
doc.add_paragraph('将 conversation-canvas.exe 复制到任意位置，双击运行即可。')

p = doc.add_paragraph()
run = p.add_run('注意：')
run.bold = True
p.add_run(' 首次运行时，Windows 可能弹出安全提示，点击"更多信息"→"仍要运行"。')

# ═══════════════════════════════════════════════════
#  3. 准备 API Key
# ═══════════════════════════════════════════════════
doc.add_heading('3. 准备 API Key', level=1)

doc.add_paragraph('使用前需要一个 AI 模型的 API Key。推荐使用 DeepSeek：')

steps = [
    '打开 https://platform.deepseek.com',
    '注册账号（新用户有免费额度）',
    '登录后进入「API Keys」页面',
    '点击「创建 API Key」，复制保存',
]
for i, step in enumerate(steps):
    doc.add_paragraph(f'{i+1}. {step}')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('支持的模型：')
run.bold = True

models = [
    'DeepSeek V4 Pro / Flash（推荐，便宜好用）',
    'GPT-4o / GPT-4o Mini',
    '通义千问 Max / Plus',
    'Moonshot v1 128K',
    'GLM-4 Plus',
    '任何 OpenAI 兼容格式的模型',
]
for m in models:
    doc.add_paragraph(m, style='List Bullet')

# ═══════════════════════════════════════════════════
#  4. 使用方法
# ═══════════════════════════════════════════════════
doc.add_heading('4. 使用方法', level=1)

doc.add_heading('4.1 配置模型', level=2)
doc.add_paragraph('打开应用后，在左上角找到「⚙ 模型配置」节点：')
steps = [
    '「模型」下拉框选择你的模型（如 DeepSeek V4 Pro）',
    '「API Key」输入框填入你的 API Key',
    '「System Prompt」可选，输入系统提示词（如"你是一个专业的编程助手"）',
]
for i, s in enumerate(steps):
    doc.add_paragraph(f'{i+1}. {s}')

doc.add_heading('4.2 普通对话', level=2)
doc.add_paragraph('默认进入对话模式，操作步骤：')
steps = [
    '在「✦ 对话输入」节点的文本框中输入你的问题',
    '点击「▸ 发送」按钮',
    'AI 回复会自动显示在新生成的「◉ AI回复」节点中',
    '回复完成后，自动创建新的输入节点，可以继续对话',
]
for i, s in enumerate(steps):
    doc.add_paragraph(f'{i+1}. {s}')

doc.add_heading('4.3 Agent 模式', level=2)
doc.add_paragraph('点工具栏的「🤖 对话」按钮切换到 Agent 模式（按钮变紫色）。')
doc.add_paragraph('Agent 模式下，AI 可以自动执行任务：')
steps = [
    '点击 📁 按钮选择工作目录（AI 在这个目录下创建文件）',
    '输入任务，如"帮我创建一个 hello.py"',
    'AI 会自动思考 → 调用工具 → 创建文件',
    '节点图上实时显示每一步：🧠 思考 → 🔧 工具调用 → 📋 结果 → 📄 文件',
    '文件创建后，点击「▶ 打开文件」或「📂 打开文件夹」查看',
]
for i, s in enumerate(steps):
    doc.add_paragraph(f'{i+1}. {s}')

doc.add_heading('4.4 分支对话', level=2)
doc.add_paragraph(
    '对话画布的核心特性是分叉——你可以从任意历史节点创建分支，'
    '测试不同的问题方向，互不干扰。'
)
doc.add_paragraph('操作方法：右键画布空白处 → 添加新节点 → 手动连线到你想分叉的节点。')

doc.add_heading('4.5 对比分支', level=2)
doc.add_paragraph(
    '选中 2 个以上的回复节点（Ctrl+点击），点工具栏的「⚡ 对比分支」，'
    '可以并排对比多个回复内容。'
)

doc.add_heading('4.6 搜索节点', level=2)
doc.add_paragraph('工具栏中间的搜索框可以搜索所有节点中的对话内容，输入关键词即可定位。')

doc.add_heading('4.7 保存和加载', level=2)
doc.add_paragraph('工具栏左侧：')
doc.add_paragraph('💾 保存 — 将当前对话图导出为 JSON 文件', style='List Bullet')
doc.add_paragraph('📂 加载 — 从 JSON 文件恢复对话图', style='List Bullet')

# ═══════════════════════════════════════════════════
#  5. 画布操作
# ═══════════════════════════════════════════════════
doc.add_heading('5. 画布操作', level=1)

table = doc.add_table(rows=6, cols=2)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

table.rows[0].cells[0].text = '操作'
table.rows[0].cells[1].text = '方法'

shortcuts = [
    ['缩放画布', '鼠标滚轮'],
    ['拖动画布', '按住鼠标中键拖拽'],
    ['选中节点', '单击节点'],
    ['多选节点', 'Ctrl + 单击'],
    ['删除节点', '选中后按 Delete'],
]
for r, (op, method) in enumerate(shortcuts):
    table.rows[r+1].cells[0].text = op
    table.rows[r+1].cells[1].text = method

# ═══════════════════════════════════════════════════
#  6. 常见问题
# ═══════════════════════════════════════════════════
doc.add_heading('6. 常见问题', level=1)

qa = [
    ('Q: 双击 exe 没反应？', 'A: 可能被杀毒软件拦截。右键 → 以管理员身份运行，或将 exe 加入白名单。'),
    ('Q: 发送消息后没回复？', 'A: 检查 API Key 是否正确，模型是否选对。可以在 https://platform.deepseek.com 确认 Key 有效。'),
    ('Q: Agent 模式报"路径越界"？', 'A: 点 📁 按钮重新选择工作目录，确保目录存在且有写入权限。'),
    ('Q: 文件创建了但找不到？', 'A: 检查工作目录设置。Agent 模式下 📁 按钮显示的路径就是工作目录。'),
    ('Q: 能用其他模型吗？', 'A: 可以，任何 OpenAI 兼容格式的 API 都行。在模型配置节点选"自定义模型"，填入 Base URL 和模型名。'),
]
for q, a in qa:
    p = doc.add_paragraph()
    run = p.add_run(q)
    run.bold = True
    doc.add_paragraph(a)

# ═══════════════════════════════════════════════════
#  7. 节点类型说明
# ═══════════════════════════════════════════════════
doc.add_heading('7. 节点类型说明', level=1)

table = doc.add_table(rows=8, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

table.rows[0].cells[0].text = '节点'
table.rows[0].cells[1].text = '颜色'
table.rows[0].cells[2].text = '说明'

nodes = [
    ['⚙ 模型配置', '红色', '选择模型、填 API Key'],
    ['✦ 对话输入', '绿色', '输入问题并发送'],
    ['◉ AI回复', '黄色', '显示 AI 的回复内容'],
    ['🧠 思考', '紫色', 'Agent 的推理过程（仅 Agent 模式）'],
    ['🔧 工具调用', '蓝色', 'AI 调用了什么工具（仅 Agent 模式）'],
    ['📋 工具结果', '绿/红色', '工具执行的结果（仅 Agent 模式）'],
    ['📄 文件产物', '蓝色', 'AI 创建的文件（仅 Agent 模式）'],
]
for r, (name, color, desc) in enumerate(nodes):
    table.rows[r+1].cells[0].text = name
    table.rows[r+1].cells[1].text = color
    table.rows[r+1].cells[2].text = desc

# ═══════════════════════════════════════════════════
#  页脚
# ═══════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('GitHub: https://github.com/kkry486/conversation-canvas')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(128, 128, 128)

# 保存
output_path = 'e:/aworld/big/对话画布使用说明.docx'
doc.save(output_path)
print(f'已保存: {output_path}')
