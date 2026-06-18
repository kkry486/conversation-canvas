from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for i in range(1, 4):
    heading_style = doc.styles[f'Heading {i}']
    heading_style.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

title = doc.add_heading('对话画布（Conversation Canvas）创新设计方案', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('—— 基于ComfyUI节点图交互的非线性对话系统')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('版本：v1.0  |  日期：2026年6月17日')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

doc.add_heading('目录', level=1)
toc_items = [
    '1. 设计动机与核心理念',
    '  1.1 现有对话产品的局限',
    '  1.2 核心设计思想',
    '  1.3 与现有方案的本质区别',
    '2. 系统架构',
    '  2.1 数据模型',
    '  2.2 上下文继承机制',
    '  2.3 布局算法',
    '3. 交互设计',
    '  3.1 默认视图：线性模式',
    '  3.2 画布视图',
    '  3.3 关键交互流程',
    '4. 技术实现方案',
    '  4.1 前端画布技术选型',
    '  4.2 后端架构',
    '  4.3 搜索与导航',
    '  4.4 数据持久化',
    '5. 竞品对比分析',
    '6. 应用场景',
    '7. 挑战与解决方案',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# === Section 1 ===
doc.add_heading('1. 设计动机与核心理念', level=1)

doc.add_heading('1.1 现有对话产品的局限', level=2)
doc.add_paragraph(
    '当前所有主流AI对话产品（ChatGPT、Claude、文心一言、Kimi等）都采用线性列表的交互形式：'
    '用户发一条消息，AI回一条消息，依次向下排列。这种结构在大多数日常场景下足够使用，'
    '但在需要深度探索、多线程思考、或反复迭代的场景中，暴露出了明显的局限性：'
)

problems = [
    ('无法分支', '聊到第8轮发现方向错了，想回到第3轮换个方向重新探索。现在只能开一个新对话，把前面的内容重新描述一遍，丢失了完整的上下文。'),
    ('无法对比', '想看"如果换一种问法会怎样"，没有便捷的方式在同一界面中并行探索多个方向并对比结果。'),
    ('无法回溯', '对话历史是一条直线，不支持从任意历史节点重新出发。用户被锁定在"只能在末尾继续"的模式中。'),
    ('话题管理困难', '在一段长对话中岔出去聊新话题后，很难再清晰地回到原来的主线。话题混杂在一起，认知负担大。'),
]
for title_text, desc in problems:
    p = doc.add_paragraph()
    run = p.add_run(f'  {title_text}：')
    run.bold = True
    p.add_run(desc)

doc.add_heading('1.2 核心设计思想', level=2)
doc.add_paragraph(
    '本方案的核心思想是：把对话本身变成一张可视化的关系图（有向无环图DAG），'
    '借鉴ComfyUI的节点编辑器交互方式，让用户可以在画布上直观地看到、创建和管理对话的分支结构。'
)
doc.add_paragraph('具体而言：')
points = [
    '每一轮对话（用户消息 + AI回复）成为画布上的一个节点',
    '节点之间用连线表示对话的线性流程',
    '用户可以从任意历史节点拉一条线，开一个新的对话分支',
    '新分支自动继承父节点的完整对话上下文',
    '对话结构在画布上一目了然，支持缩放、拖拽、搜索、对比',
]
for point in points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_heading('1.3 与现有方案的本质区别', level=2)

table = doc.add_table(rows=7, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['维度', '现有对话产品', '对话画布设计']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

data = [
    ['对话结构', '线性列表', '有向图（DAG）'],
    ['继续对话', '只能在末尾', '从任意历史节点分叉'],
    ['分支可见性', '不可见/折叠', '画布上清晰可见'],
    ['上下文处理', '不同分支互不影响', '新分支继承父节点上下文'],
    ['信息呈现', '扁平的文字排列', '空间化的节点布局'],
    ['类比', '记事本', '对话的Git——分支、合并、回溯'],
]
for row_idx, row_data in enumerate(data):
    for col_idx, cell_text in enumerate(row_data):
        table.rows[row_idx + 1].cells[col_idx].text = cell_text

# === Section 2 ===
doc.add_page_break()
doc.add_heading('2. 系统架构', level=1)

doc.add_heading('2.1 数据模型', level=2)
doc.add_paragraph('对话画布的核心数据结构包含两个部分：对话节点（ConversationNode）和对话图（ConversationGraph）。')

doc.add_paragraph('对话节点（ConversationNode）的字段定义：')
node_fields = [
    ('id', '字符串，唯一标识'),
    ('parentId', '字符串或null，父节点ID，根节点为null'),
    ('branchIndex', '整数，在父节点下的第几个分支'),
    ('userMessage', '字符串，用户消息内容'),
    ('aiResponse', '字符串，AI回复内容'),
    ('timestamp', '时间戳，创建时间'),
    ('context', '消息数组，该节点的完整上下文（含祖先链路）'),
    ('position', '坐标对象，画布上的x/y位置'),
    ('metadata', '元数据对象，包含token消耗、模型信息、工具调用记录等'),
]
table2 = doc.add_table(rows=len(node_fields)+1, cols=2)
table2.style = 'Light Grid Accent 1'
table2.rows[0].cells[0].text = '字段'
table2.rows[0].cells[1].text = '说明'
for p in table2.rows[0].cells[0].paragraphs:
    for r in p.runs:
        r.bold = True
for p in table2.rows[0].cells[1].paragraphs:
    for r in p.runs:
        r.bold = True
for i, (field, desc) in enumerate(node_fields):
    table2.rows[i+1].cells[0].text = field
    table2.rows[i+1].cells[1].text = desc

doc.add_paragraph()
doc.add_paragraph('对话图（ConversationGraph）的字段定义：')
graph_fields = [
    ('rootId', '字符串，根节点ID'),
    ('nodes', 'Map<String, ConversationNode>，所有节点的映射'),
    ('branches', 'Map<String, String[]>，parentId到子节点ID列表的映射'),
]
table3 = doc.add_table(rows=len(graph_fields)+1, cols=2)
table3.style = 'Light Grid Accent 1'
table3.rows[0].cells[0].text = '字段'
table3.rows[0].cells[1].text = '说明'
for p in table3.rows[0].cells[0].paragraphs:
    for r in p.runs:
        r.bold = True
for p in table3.rows[0].cells[1].paragraphs:
    for r in p.runs:
        r.bold = True
for i, (field, desc) in enumerate(graph_fields):
    table3.rows[i+1].cells[0].text = field
    table3.rows[i+1].cells[1].text = desc

doc.add_heading('2.2 上下文继承机制', level=2)
doc.add_paragraph(
    '上下文继承是对话画布的核心机制之一。当用户从节点A（比如第3轮对话）分叉出新分支时，'
    '新分支的上下文自动组装为：从根节点到节点A的完整对话历史 + 新的用户消息。'
    '这意味着AI能够"记得"之前的讨论内容，无需用户重新描述背景。'
)
doc.add_paragraph('具体的上下文组装规则：')
rules = [
    '新分支的上下文 = 父节点的context + 新的用户消息',
    '如果上下文总长度超过模型的context window，触发自动压缩机制',
    '压缩策略：保留最近N轮完整对话，更早的历史用摘要替代',
    '用户可手动选择上下文深度："带多少轮历史"',
]
for rule in rules:
    doc.add_paragraph(rule, style='List Bullet')

doc.add_heading('2.3 布局算法', level=2)
doc.add_paragraph(
    '画布上的节点布局采用层次化自动布局算法（类似Graphviz的DOT算法），确保节点排列清晰有序：'
)
layout_rules = [
    '根节点放在画布最左侧',
    '同一对话深度的节点垂直排列',
    '分支向右延伸，层级递进',
    '节点间距自适应内容长度',
    '连线自动避障，尽量不交叉',
    '支持用户手动拖拽调整节点位置，手动布局优先于自动布局',
]
for rule in layout_rules:
    doc.add_paragraph(rule, style='List Bullet')

# === Section 3 ===
doc.add_page_break()
doc.add_heading('3. 交互设计', level=1)

doc.add_heading('3.1 默认视图：线性模式', level=2)
doc.add_paragraph(
    '为了降低用户的学习成本，产品默认提供线性对话视图，体验与现有AI对话产品完全一致。'
    '只有在用户主动触发时，才展开画布视图。这样既保证了普通用户的使用体验，'
    '又为需要高级功能的用户提供了强大的工具。'
)
doc.add_paragraph('触发画布视图的方式：')
triggers = [
    '按钮触发：界面上提供"切换到画布视图"按钮',
    '右键菜单：在某个历史消息上右键 → "从此处分叉"',
    '手势操作：在移动端，长按某个历史消息触发分支选项',
    '快捷键：Ctrl+B（Branch）快速创建分支',
]
for trigger in triggers:
    doc.add_paragraph(trigger, style='List Bullet')

doc.add_heading('3.2 画布视图', level=2)
doc.add_paragraph('进入画布视图后，用户看到的是一个类似ComfyUI的节点编辑器界面，核心元素包括：')
canvas_elements = [
    ('节点卡片', '每个对话轮次显示为一个卡片，包含用户消息摘要和AI回复摘要。卡片大小自适应内容长度。'),
    ('连线', '从父节点指向子节点的有向箭头，不同分支用不同颜色区分，当前激活路径高亮显示。'),
    ('缩放与拖拽', '支持画布无限缩放和平移，节点支持手动拖拽调整位置。'),
    ('节点详情', '点击节点卡片展开完整的对话内容，包括AI的思考过程、调用的工具等。'),
    ('分支操作', '从任意节点的边缘拉线创建新分支，或通过右键菜单操作。'),
    ('全局缩略图', '画布右下角显示全局缩略图，支持快速定位到画布的任意区域。'),
]
for name, desc in canvas_elements:
    p = doc.add_paragraph()
    run = p.add_run(f'  {name}：')
    run.bold = True
    p.add_run(desc)

doc.add_heading('3.3 关键交互流程', level=2)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('创建分支的流程：')
run.bold = True
steps = [
    '用户在画布上点击某个历史节点',
    '节点周围出现"+"按钮或"从此处分叉"选项',
    '用户输入新的消息（在弹出的输入框中）',
    '画布上生成新节点，自动连线到父节点',
    'AI基于父节点的上下文生成回复',
    '新节点出现在画布上，连线颜色与父节点所在分支一致',
]
for i, step in enumerate(steps):
    doc.add_paragraph(f'  第{i+1}步：{step}')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('切换分支的流程：')
run.bold = True
doc.add_paragraph('  点击任意分支的任意节点 → 该分支路径高亮 → 可以继续在该分支上对话')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('对比视图的流程：')
run.bold = True
doc.add_paragraph('  选中两个不同分支的节点 → 并排显示两个节点的AI回复 → 差异高亮显示')

# === Section 4 ===
doc.add_page_break()
doc.add_heading('4. 技术实现方案', level=1)

doc.add_heading('4.1 前端画布技术选型', level=2)

table4 = doc.add_table(rows=5, cols=3)
table4.style = 'Light Grid Accent 1'
table4.rows[0].cells[0].text = '技术'
table4.rows[0].cells[1].text = '特点'
table4.rows[0].cells[2].text = '适用场景'
for p in table4.rows[0].cells[0].paragraphs:
    for r in p.runs:
        r.bold = True
for p in table4.rows[0].cells[1].paragraphs:
    for r in p.runs:
        r.bold = True
for p in table4.rows[0].cells[2].paragraphs:
    for r in p.runs:
        r.bold = True

tech_data = [
    ['React Flow', '最流行的节点图React库，生态丰富，社区活跃', '首选方案，适合快速开发'],
    ['AntV X6', '蚂蚁集团图编辑引擎，功能强大，支持复杂交互', '需要高度定制化交互时'],
    ['LiteGraph.js', '轻量级，ComfyUI同款引擎', '追求极致性能和轻量化时'],
    ['D3.js', '灵活的底层可视化库', '需要完全自定义渲染逻辑时'],
]
for i, row_data in enumerate(tech_data):
    for j, cell_text in enumerate(row_data):
        table4.rows[i+1].cells[j].text = cell_text

doc.add_heading('4.2 后端架构', level=2)
doc.add_paragraph('后端系统由四个核心层组成：')

layers = [
    ('API Gateway', '统一的API入口，负责请求路由、认证、限流。'),
    ('Conversation Service', '对话服务层，包含四个核心模块：'),
    ('AI Service', 'AI服务层，负责模型调用、检索增强、工具执行。'),
    ('Storage', '存储层，包含图数据库、向量数据库和缓存。'),
]
for name, desc in layers:
    p = doc.add_paragraph()
    run = p.add_run(f'  {name}：')
    run.bold = True
    p.add_run(desc)

doc.add_paragraph('Conversation Service的四个核心模块：')
modules = [
    'GraphManager：管理对话图的节点增删改查，维护节点间的父子关系',
    'ContextBuilder：根据节点的祖先链路组装完整的对话上下文，处理上下文压缩',
    'LayoutEngine：计算节点在画布上的最优布局坐标',
    'SyncService：处理多端同步和多人协作的实时状态同步',
]
for module in modules:
    doc.add_paragraph(module, style='List Bullet')

doc.add_heading('4.3 搜索与导航', level=2)
doc.add_paragraph('为了帮助用户在复杂的对话图中快速定位，系统提供以下搜索与导航功能：')
search_features = [
    '关键词搜索：支持按关键词搜索历史对话节点内容',
    '语义搜索：基于向量化技术，支持语义层面的对话检索',
    '路径高亮：从根节点到当前节点的完整路径高亮显示',
    '面包屑导航：显示当前节点的祖先链路，支持快速跳转',
    '缩略图导航：画布右下角显示全局缩略图，支持快速定位',
    '聚焦模式：只显示从根节点到当前节点的路径，隐藏其他分支',
]
for feature in search_features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_heading('4.4 数据持久化', level=2)

table5 = doc.add_table(rows=5, cols=3)
table5.style = 'Light Grid Accent 1'
table5.rows[0].cells[0].text = '存储类型'
table5.rows[0].cells[1].text = '技术选型'
table5.rows[0].cells[2].text = '存储内容'
for p in table5.rows[0].cells[0].paragraphs:
    for r in p.runs:
        r.bold = True
for p in table5.rows[0].cells[1].paragraphs:
    for r in p.runs:
        r.bold = True
for p in table5.rows[0].cells[2].paragraphs:
    for r in p.runs:
        r.bold = True

storage_data = [
    ['图数据库', 'Neo4j / ArangoDB', '对话图的节点关系、分支结构'],
    ['文档数据库', 'MongoDB / PostgreSQL', '对话内容、元数据、用户信息'],
    ['向量数据库', 'Pinecone / Weaviate / Milvus', '向量化的对话片段，支持语义搜索'],
    ['缓存', 'Redis', '活跃对话图结构的热数据缓存，加速访问'],
]
for i, row_data in enumerate(storage_data):
    for j, cell_text in enumerate(row_data):
        table5.rows[i+1].cells[j].text = cell_text

# === Section 5 ===
doc.add_page_break()
doc.add_heading('5. 竞品对比分析', level=1)

doc.add_paragraph(
    '目前市场上没有与本方案完全对标的竞品。以下对比最接近的相关产品和功能：'
)

table6 = doc.add_table(rows=6, cols=5)
table6.style = 'Light Grid Accent 1'
headers6 = ['产品/功能', '对话分支', '可视化画布', '上下文继承', '节点图交互']
for i, h in enumerate(headers6):
    table6.rows[0].cells[i].text = h
    for p in table6.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

comp_data = [
    ['ChatGPT Branch', '有限支持', '无', '无', '无'],
    ['Claude 对话', '不支持', '无', '无', '无'],
    ['对话画布（本设计）', '任意节点分叉', 'ComfyUI风格', '完整继承', '全功能节点图'],
    ['LangGraph Studio', 'N/A（面向工作流）', '有', 'N/A', '有，但面向开发者'],
    ['思维导图工具', '手动创建', '有', '无', '有，但非AI对话'],
]
for i, row_data in enumerate(comp_data):
    for j, cell_text in enumerate(row_data):
        table6.rows[i+1].cells[j].text = cell_text

doc.add_paragraph()
doc.add_paragraph(
    '从对比中可以看出，本设计填补了"对话级可视化节点图"这一市场空白。'
    'ChatGPT的Branch功能虽然支持分支，但完全没有可视化，且分支后旧内容被折叠；'
    'LangGraph Studio有节点图，但面向的是工作流编排而非用户对话；'
    '思维导图工具有可视化，但不是AI对话系统。'
    '本方案将这三者的需求融合在了一起。'
)

# === Section 6 ===
doc.add_page_break()
doc.add_heading('6. 应用场景', level=1)

scenarios = [
    ('场景一：研究探索',
     '研究者在做一个课题，对话到第8轮发现方向有偏差，想回到第3轮换个角度重新探索。'
     '直接从第3轮开新分支，新对话自动带着前3轮的背景知识，无需重新描述上下文。'
     '多个研究方向在画布上并行展开，一目了然。'),
    ('场景二：版本对比',
     '让AI帮你写一段文案，写到第5轮不满意。从第3轮开新分支，用不同的方式重新提问，'
     '在画布上并排比较两个分支的产出，选择更好的方向继续深入。'
     '就像代码的Git分支一样，探索不同的"版本"。'),
    ('场景三：思维发散',
     '在头脑风暴中，一个话题聊到深处想岔出去问一个相关但不同方向的问题。'
     '分支出去不影响主线，随时可以切回来继续原来的讨论。'
     '画布上清晰展示思维的发散和收敛过程。'),
    ('场景四：技术方案讨论',
     '讨论一个技术方案时，主分支走方案A。在某个关键决策点分出方案B、方案C，'
     '在画布上一目了然地对比不同方案的讨论过程和最终结果。'
     '决策点和各方案的优劣清晰可见。'),
    ('场景五：教学与培训',
     '讲师和AI对话演示某个概念，从不同角度分出多个分支讲解不同方面。'
     '学生可以在画布上看到整个知识体系的展开过程，'
     '理解各个概念之间的关联和层级关系。'),
    ('场景六：创意写作',
     '写小说或剧本时，从同一个故事节点分出不同的剧情走向，'
     '在画布上同时探索多条故事线，比较不同走向的戏剧效果，'
     '最终选择最满意的路径继续创作。'),
]
for title_text, desc in scenarios:
    p = doc.add_paragraph()
    run = p.add_run(title_text)
    run.bold = True
    doc.add_paragraph(desc)

# === Section 7 ===
doc.add_page_break()
doc.add_heading('7. 挑战与解决方案', level=1)

challenges = [
    ('挑战一：画布复杂度管理',
     '问题：分支多了之后节点密密麻麻，用户看不清整体结构。',
     [
         '默认折叠非活跃分支，只展开当前查看的链路',
         '提供"聚焦模式"：只显示从根节点到当前节点的路径',
         '支持按分支或主题分组折叠',
         '智能布局算法自动整理节点位置，避免拥挤',
         '支持搜索定位，输入关键词快速跳转到目标节点',
     ]),
    ('挑战二：上下文长度限制',
     '问题：深层嵌套分支的上下文可能超过模型的context window（如128K token）。',
     [
         '自动摘要机制：对祖先节点的历史对话做摘要压缩',
         '滑动窗口策略：保留最近N轮完整对话，更早的用摘要替代',
         'RAG辅助：将完整历史向量化存储，需要时检索相关片段',
         '用户可手动选择上下文深度："带多少轮历史"',
         '智能上下文选择：根据新问题的语义相关性，自动选取最相关的祖先对话',
     ]),
    ('挑战三：状态同步与协作',
     '问题：多设备访问或多人协作时，图结构的一致性维护。',
     [
         'CRDT（无冲突复制数据类型）实现协作编辑，避免冲突',
         'WebSocket实时同步节点变更，保证多端一致',
         '版本控制：每次变更生成快照，支持回滚到任意历史状态',
         '操作日志：记录所有分支创建、修改操作，便于审计和追溯',
     ]),
    ('挑战四：用户学习成本',
     '问题：非技术用户可能不熟悉节点图交互，感觉复杂。',
     [
         '默认线性视图，零学习成本，与现有产品体验一致',
         '画布视图作为"高级模式"渐进开放，不强制使用',
         '交互引导：首次使用画布时通过动画演示如何创建分支',
         '预设模板：提供常见的分支模式（如"探索-对比-选择"模板）',
         '渐进式功能暴露：先展示基础的分叉功能，高级功能（对比视图、聚焦模式等）在用户熟悉后再解锁',
     ]),
    ('挑战五：性能优化',
     '问题：大型对话图（数百个节点）的渲染和交互性能。',
     [
         '虚拟化渲染：只渲染可视区域内的节点，超出区域的懒加载',
         '图结构索引：使用空间索引（如R-tree）加速节点查找和碰撞检测',
         '增量更新：只更新变化的节点和连线，不重绘整个画布',
         '服务端预计算布局：复杂图的布局在服务端计算好再下发，减轻前端负担',
     ]),
]
for title_text, problem, solutions in challenges:
    p = doc.add_paragraph()
    run = p.add_run(title_text)
    run.bold = True
    run.font.size = Pt(12)
    doc.add_paragraph(problem)
    p2 = doc.add_paragraph()
    run2 = p2.add_run('解决方案：')
    run2.bold = True
    for solution in solutions:
        doc.add_paragraph(solution, style='List Bullet')
    doc.add_paragraph()

# === 结语 ===
doc.add_page_break()
doc.add_heading('结语', level=1)
doc.add_paragraph(
    '对话画布（Conversation Canvas）的设计方案，本质上是将版本控制的分支思想引入AI对话界面，'
    '并用ComfyUI的节点图交互方式呈现出来。它解决的核心问题是：现有AI对话产品的线性结构'
    '限制了用户进行深度探索、多线程思考和迭代对比的能力。'
)
doc.add_paragraph(
    '通过将每轮对话可视化为节点、支持任意历史节点的分支创建、以及自动化的上下文继承，'
    '对话画布让用户能够在一张画布上看到完整的思维脉络，自由地探索不同的对话方向，'
    '并在多个分支之间进行直观的对比。'
)
doc.add_paragraph(
    '这个方向目前在市面上没有直接的对标产品，填补了"对话级可视化节点图"的市场空白。'
    '最大的挑战在于如何让画布在分支变多时依然清晰可读，以及如何平衡"可视化的力量"'
    '和"界面的简洁性"。本方案通过"默认线性 + 渐进式画布"的策略，'
    '试图在这两者之间找到最佳平衡点。'
)

# Save
output_path = r'C:\Users\空空如也\AppData\Local\Claude-3p\local-agent-mode-sessions\af3c8899\00000000\local_fea3584a-2c4b-42e8-93b3-9de8374f064a\outputs\Conversation_Canvas_对话画布设计方案.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
