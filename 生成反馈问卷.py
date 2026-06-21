"""生成对话画布反馈问卷 Word 文档"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.6

# ═══════════════════════════════════════════════════
#  标题
# ═══════════════════════════════════════════════════
title = doc.add_heading('对话画布 — 用户反馈问卷', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('感谢你花几分钟试用这个工具，你的反馈会直接决定这个项目下一步怎么做。')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

# ═══════════════════════════════════════════════════
#  第一部分：基本信息
# ═══════════════════════════════════════════════════
doc.add_heading('一、基本信息', level=1)

questions = [
    ('1. 你的技术背景是？', '□ 开发者（写过代码）\n□ 非开发者（产品/设计/运营等）\n□ 学生（计算机相关）\n□ 学生（非计算机）'),
    ('2. 你平时用 AI 工具吗？', '□ 每天都用\n□ 偶尔用\n□ 很少用\n□ 从没用过'),
    ('3. 你用过哪些 AI 工具？（多选）', '□ ChatGPT / Claude 网页版\n□ Cursor / Windsurf\n□ Claude Code\n□ GitHub Copilot\n□ 其他：________'),
]

for q, options in questions:
    p = doc.add_paragraph()
    run = p.add_run(q)
    run.bold = True
    doc.add_paragraph(options)
    doc.add_paragraph()

# ═══════════════════════════════════════════════════
#  第二部分：安装体验
# ═══════════════════════════════════════════════════
doc.add_heading('二、安装体验', level=1)

questions = [
    ('4. 你拿到 exe 后多久跑起来的？', '□ 1 分钟内\n□ 1-5 分钟\n□ 5-15 分钟\n□ 超过 15 分钟 / 没跑起来'),
    ('5. 安装过程中遇到问题了吗？', '□ 没有，很顺利\n□ 有，具体是：________________________________\n□ 完全跑不起来'),
    ('6. 使用说明文档有帮助吗？', '□ 有帮助，看懂了\n□ 有帮助，但有些地方不清楚\n□ 没看，直接用的\n□ 看了但没看懂'),
]

for q, options in questions:
    p = doc.add_paragraph()
    run = p.add_run(q)
    run.bold = True
    doc.add_paragraph(options)
    doc.add_paragraph()

# ═══════════════════════════════════════════════════
#  第三部分：核心体验
# ═══════════════════════════════════════════════════
doc.add_heading('三、核心体验', level=1)

questions = [
    ('7. 节点图画布的交互体验如何？', '□ 很好，直观易懂\n□ 还行，但需要适应\n□ 有点复杂，不太好操作\n□ 完全不会用'),
    ('8. AI 回复的质量如何？', '□ 很好，回答准确\n□ 还行，基本能用\n□ 一般，经常答非所问\n□ 很差，基本没用'),
    ('9. 你用了对话模式还是 Agent 模式？', '□ 只用了对话模式\n□ 只用了 Agent 模式\n□ 两个都用了\n□ 不知道有 Agent 模式'),
]

for q, options in questions:
    p = doc.add_paragraph()
    run = p.add_run(q)
    run.bold = True
    doc.add_paragraph(options)
    doc.add_paragraph()

# ═══════════════════════════════════════════════════
#  第四部分：Agent 模式专项
# ═══════════════════════════════════════════════════
doc.add_heading('四、Agent 模式专项（如果用了的话）', level=1)

questions = [
    ('10. Agent 创建文件成功了吗？', '□ 成功了，能看到文件产物节点\n□ 失败了，报错：________________\n□ 没试这个功能'),
    ('11. 节点图展示 Agent 的思考过程，你觉得有用吗？', '□ 很有用，能看懂 AI 在做什么\n□ 有点用，但信息太多了\n□ 没什么用，看不过来\n□ 不知道有这个功能'),
    ('12. 和直接用 ChatGPT/Claude 对话比，节点图方式有什么优势？（开放题）', '________________________________\n\n________________________________'),
    ('13. 和直接用 ChatGPT/Claude 对话比，节点图方式有什么劣势？（开放题）', '________________________________\n\n________________________________'),
]

for q, options in questions:
    p = doc.add_paragraph()
    run = p.add_run(q)
    run.bold = True
    doc.add_paragraph(options)
    doc.add_paragraph()

# ═══════════════════════════════════════════════════
#  第五部分：需求验证
# ═══════════════════════════════════════════════════
doc.add_heading('五、需求验证（最重要）', level=1)

p = doc.add_paragraph()
run = p.add_run('以下问题决定项目是否值得继续做，请认真回答。')
run.font.color.rgb = RGBColor(200, 50, 50)

doc.add_paragraph()

questions = [
    ('14. 你会继续用这个工具吗？', '□ 会，已经替代了我之前的 AI 工具\n□ 可能会，看后续改进\n□ 不会，还是用回 ChatGPT/Claude\n□ 不确定'),
    ('15. 你会推荐给朋友用吗？', '□ 会，已经推荐了\n□ 可能会\n□ 不会\n□ 看情况'),
    ('16. 如果这个工具收费，你愿意付多少钱？', '□ 免费才用\n□ 10 元/月以内\n□ 10-30 元/月\n□ 30 元以上/月'),
    ('17. 你觉得最有价值的功能是？（多选）', '□ 节点图可视化对话\n□ 分叉对话\n□ 分支对比\n□ Agent 自动执行任务\n□ 多模型支持\n□ 保存/加载对话'),
    ('18. 你觉得最需要改进的是？（多选，最多选 3 个）', '□ 安装太麻烦\n□ 功能不够多\n□ 不够稳定\n□ 节点操作不直观\n□ Agent 不够聪明\n□ 其他：________'),
]

for q, options in questions:
    p = doc.add_paragraph()
    run = p.add_run(q)
    run.bold = True
    doc.add_paragraph(options)
    doc.add_paragraph()

# ═══════════════════════════════════════════════════
#  第六部分：开放反馈
# ═══════════════════════════════════════════════════
doc.add_heading('六、开放反馈', level=1)

questions = [
    ('19. 这个工具最让你印象深刻的是什么？', '________________________________\n\n________________________________'),
    ('20. 最让你失望的是什么？', '________________________________\n\n________________________________'),
    ('21. 你希望它增加什么功能？', '________________________________\n\n________________________________'),
    ('22. 其他想说的：', '________________________________\n\n________________________________'),
]

for q, options in questions:
    p = doc.add_paragraph()
    run = p.add_run(q)
    run.bold = True
    doc.add_paragraph(options)
    doc.add_paragraph()

# ═══════════════════════════════════════════════════
#  页脚
# ═══════════════════════════════════════════════════
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('感谢你的反馈！每一份问卷都会认真阅读。')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(100, 100, 100)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('联系方式：3132834400@qq.com')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(150, 150, 150)

# 保存
output_path = 'e:/aworld/big/对话画布反馈问卷.docx'
doc.save(output_path)
print(f'已保存: {output_path}')
